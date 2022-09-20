# -*- coding: utf-8 -*-
"""
Invoice Automation

Required Pakages:

    pdfplumber - pip install pdfplumber
    pdfquery - pip install pdfquery
    docx - pip install python-docx

@author: bhargav
"""

import pdfquery
import pdfplumber
import docx
import re
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT


def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def scrape(key, tol=0):
    keyword =  pdf.pq('LTTextLineHorizontal:contains("{}")'.format(key))[0]
    x0 = float(keyword.get('x0',0))
    y0 = float(keyword.get('y0',0)) - tol
    x1 = float(keyword.get('x1',0))
    y1 = float(keyword.get('y1',0)) - tol
    res = pdf.pq('LTTextLineHorizontal:overlaps_bbox("%s, %s, %s, %s")' % (x0, y0, x1, y1)).text()
    return(res)

def scrape2(key, tol=0):
    keyword =  pdf.pq('LTTextLineHorizontal:contains("{}")'.format(key))[0]
    x0 = float(keyword.get('x0',0)) #-tol
    y0 = float(keyword.get('y0',0)) - tol
    x1 = float(keyword.get('x1',0)) #+ tol
    y1 = float(keyword.get('y1',0)) #+ tol
    res = pdf.pq('LTTextLineHorizontal:overlaps_bbox("%s, %s, %s, %s")' % (x0, y0, x1, y1)).text()
    #print(x0, y0, x1, y1)
    return(res)

pdf = ""

def automate(inFile,outFile, product_id, invoice_id):
    global pdf
    pdf = pdfquery.PDFQuery(inFile)
    pdf.load()

    #shipTo = scrape("SHIP TO",10)
    #shipTo = shipTo[shipTo.find("SHIP TO")+9:shipTo.find("ITEM TYPE:")]

    add = scrape2("SEND INVOICES C/O – VENDOR SITE",60)

    x = add.find(".com")
    xs = add.find("SEND INVOICES C/O – VENDOR SITE")+32

    if x!=-1:
        add = add[xs:x+4]
    else:
        add = add[xs:]

    billTo = add
    req = scrape("REQUESTOR")[11:]
    eml = scrape("EMAIL")[7:]
    PO = scrape("Purchase Order No",5).split(" ")[-1]
    TIA = scrape("TOTAL AMOUNT",5)[13:]     #total invoice
    project = scrape("PROJECT",5)[30:]
    project = project[:project.find("ADDITIONAL")]
    cur = scrape("CURRENCY",5).split(" ")[-1]

    cr=""
    if cur=="USD":
        cr = "$ "
    elif cur =="GBP":
        cr = "£ "


    t=""

    with pdfplumber.open(inFile) as pdf2:
        pages = pdf2.pages
        for i,pg in enumerate(pages):
            tbl = pages[i].extract_text()
            t+=tbl
    #print(t)
    f=1
    l=[]

    while(f):
        x = t.find("QUANTITY AMOUNT",f)
        if x!=-1:
            l.append(t[x:x+50])
            f=x+50

        else:
            break

    imTyp = []
    skipping_indices = []
    for i in range(len(l)):
        x = t.find(str(i+1)+" ORC")
        if x == -1:
            skipping_indices.append(i)
            x = t.find(str(i+2)+" ORC")
            if x == -1:
                imTyp.append('')

        y = t.find("ADDITIONAL INFO/ISBN:",x)
        nx = t[x:y]

        yy = (re.search("(\d)(?!.*\d)", nx).start())+1

        xx = (re.search("\s\d{2}[-][A-z]{3}[-]\d{2}\s", nx).start())
        nx = nx.replace(nx[xx:yy],"")

        tx = nx[-2:-1]
        if ord(tx) in range(ord("0"), ord("9")+1):
            imTyp.append(" ".join(nx.split(" ")[2:7]))
        else:imTyp.append(" ".join(nx.split(" ")[2:]))

    k=[]
    increment = 1

    for i in range(len(l)):
        if i in skipping_indices:
            increment += 1
        tmp = l[i].split(" ")
        tmp[6] = tmp[6].replace(',', '')
        k.append([tmp[3], str(i + increment), imTyp[i], tmp[5], cr +
            ("{:.2f}".format(float(tmp[6])/float(tmp[5]))), cr+tmp[6]])

    k.append(["","","","","Subtotal",cr+str(TIA)])
    k.append(["","","","","Tax",cr+str(0.00)])
    k.append(["","","","","*Total Due",cr+str(TIA)])



    # Create an instance of a word document
    doc = docx.Document()

    style = doc.styles['Heading 4']
    font = style.font
    font.size = Pt(36)
    font.italic = True
    font.bold = False
    font.name = 'Times New Roman'

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Verdana'
    font.size = Pt(10)

    # Add a Title to the document
    tb = doc.add_paragraph("INVOICE")
    tb.alignment = 1
    tb.style = doc.styles['Heading 4']


    ta = doc.add_paragraph("")
    ta.alignment = 1
    ta.style = doc.styles['Normal']
    p = doc.add_paragraph()
    runner = p.add_run("""Supplier Name: Continual Engine US LLC \nSupplier Address: 5440 Harvest Hill Road \nSuite 249, Dallas, Texas, 75230 \nFederal Tax ID:""")
    runner.bold = True
    p.add_run("""(if applicable)""")

    runner = p.add_run("""\nSupplier Contact:\n""")
    runner.bold = True


    p.add_run("""Vijayshree Vethantham (vijayshree.vethantham@continualengine.com)
    Accounts Receivable (accounts.receivable@continualengine.com)""")

    runner = p.add_run("""\n\nInvoice to: \n""")
    runner.bold = True

    p.add_run(billTo+"\n")
    #p.add_run(shipTo)

    runner = p.add_run("""\n\n*Attention: """)
    runner.bold = True

    p.add_run(req+"("+eml+")")

    # Table data in a form of list
    data = k

    table = doc.add_table(rows=0, cols=4, style="Table Grid")
    table.allow_autofit = True

    row = table.add_row().cells
    row[0].text = "\n*PO Number\n"
    row[1].text = "\n"+str(PO)+"\n"
    row[2].text = "\n*Invoice Number\n"
    #row[3].text = ""
    row[3].text = invoice_id
    row = table.add_row().cells
    row[0].text = "\nDelivery Date:\n"
    row[1].text = ""
    row[2].text = "\n*Invoice Date\n"
    row[3].text = ""
    row = table.add_row().cells
    row[0].text = "\nProject Name\n"
    #row[1].text = project
    row[1].text = product_id
    row[2].text = "\n*Total Invoice Amount:\n"
    row[3].text = cr+str(TIA)
    row = table.add_row().cells
    row[0].text = "Task Description"
    row[1].text = """Creation, authoring, and reviewing of alternate text (alt text) descriptions
    for """+project
    row[2].text = "*Currency:"
    row[3].text = cur

    make_rows_bold(table.rows[-1], table.rows[-2])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=6, style="Light Grid Accent 1")
    table.allow_autofit = True

    row = table.rows[0].cells
    row[0].text = '*Shipment #'
    row[1].text = '*Line Item #'
    row[2].text = '*ImageCategory'
    row[3].text = '*Quantity'
    if cr!="":
        row[4].text = '*Rate (%s)' % cr[0]
        row[5].text = '*Amount (%s)' % cr[0]
    else:
        row[4].text = '*Rate'
        row[5].text = '*Amount'

    for S, L, I, Q, R, A  in data:

        row = table.add_row().cells
        row[0].text = S
        row[1].text = L
        row[2].text = I
        row[3].text = Q
        row[4].text = R
        row[5].text = A

    make_rows_bold(table.rows[-1])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    doc.add_paragraph("")

    table = doc.add_table(rows=0, cols=2, style="Table Grid")
    table.allow_autofit = True
    row = table.add_row().cells
    row[0].text = "REMIT TO:"
    row[1].text = ""

    row = table.add_row().cells
    row[0].text = "\nAccount Owner Name:\n"
    row[1].text = "\nContinual Engine US, LLC\n"

    row = table.add_row().cells
    row[0].text = "\nPrimary Bank Name:\n"
    row[1].text = "\nJPMorgan Chase Bank, N.A \n"

    row = table.add_row().cells
    row[0].text = "\nFull Bank Address including Country:\n"
    row[1].text = "\n PO Box 182051 Columbus, OH 43218 - 2051 \n"

    row = table.add_row().cells
    row[0].text = "\nRouting Number or Sort Code:\n"
    row[1].text = "\n111000614\n"

    row = table.add_row().cells
    row[0].text = "\nBank Account Number:\n"
    row[1].text = "\n936777387\n"

    make_rows_bold(table.rows[1], table.rows[2], table.rows[3], table.rows[0], table.rows[4], table.rows[5])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


    tu = doc.add_paragraph()
    runner = tu.add_run("""\nThank you for your business!""")
    runner.bold = True
    tu.alignment = 1


    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    print(outFile)
    doc.save(outFile)

if __name__ == "__main__":
    pass
